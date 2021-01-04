import sys
import requests
from docx import Document
from docx.shared import Inches
import datetime

def to_integer(dt_time):
    return 10000*dt_time.year + 100*dt_time.month + dt_time.day


def main():
    print('Notícias observador ' + str(datetime.date.today()))

    obs_news = requests.get('https://api.observador.pt/wp/lists/featured')
    obs_news = obs_news.json()

    document1 = Document()
    document1.add_heading('Notícias Observador ' + str(datetime.date.today()))

    i = 0

    for new in obs_news:
        if (new['type'] == 'article' and  to_integer(datetime.datetime.strptime(new['pubDate'], "%Y-%m-%dT%H:%M:%S%z").date()) > (to_integer(datetime.date.today()) - 1) and i<14):
            document1.add_heading(new['title'], level=2)
            p = document1.add_paragraph()
            p.add_run(new['fullTitle']).bold = True
            p.add_run(' -> ')
            p.add_run(new['lead'])
            i = i + 1
        
    document1.save('news.docx')

if __name__ == '__main__':
    main()